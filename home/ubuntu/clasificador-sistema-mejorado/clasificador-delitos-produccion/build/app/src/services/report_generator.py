import os
import json
from datetime import datetime, timedelta
from collections import Counter, defaultdict
from typing import Dict, List, Any, Optional
import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from io import BytesIO
import base64

class ReportGenerator:
    def __init__(self, db_session):
        self.db = db_session
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Configurar estilos personalizados para el informe"""
        # Título principal
        self.styles.add(ParagraphStyle(
            name='TituloInforme',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        ))
        
        # Subtítulo
        self.styles.add(ParagraphStyle(
            name='Subtitulo',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        ))
        
        # Texto normal justificado
        self.styles.add(ParagraphStyle(
            name='TextoJustificado',
            parent=self.styles['Normal'],
            alignment=TA_JUSTIFY,
            spaceAfter=12
        ))
    
    def generate_situational_report(self, 
                                  partido: str, 
                                  fecha_inicio: str, 
                                  fecha_fin: str, 
                                  formato: str = 'pdf') -> bytes:
        """
        Generar informe situacional basado en los datos clasificados
        """
        # Obtener datos de la base de datos
        data = self._get_report_data(partido, fecha_inicio, fecha_fin)
        
        if formato.lower() == 'pdf':
            return self._generate_pdf_report(data, partido, fecha_inicio, fecha_fin)
        elif formato.lower() == 'word':
            return self._generate_word_report(data, partido, fecha_inicio, fecha_fin)
        else:
            raise ValueError(f"Formato no soportado: {formato}")
    
    def _get_report_data(self, partido: str, fecha_inicio: str, fecha_fin: str) -> Dict[str, Any]:
        """
        Obtener y procesar datos de la base de datos para el informe
        """
        # Simular datos para demostración
        # En implementación real, esto consultaría la base de datos
        
        total_delitos = 754 if partido == "José C. Paz" else 577 if partido == "San Miguel" else 675 if partido == "Malvinas Argentinas" else 1201
        
        data = {
            'partido': partido,
            'fecha_inicio': fecha_inicio,
            'fecha_fin': fecha_fin,
            'periodo': self._format_period(fecha_inicio, fecha_fin),
            'total_delitos': total_delitos,
            'delito_principal': self._get_main_crime(partido),
            'estadisticas_generales': self._generate_general_stats(total_delitos, partido),
            'distribucion_delitos': self._generate_crime_distribution(total_delitos, partido),
            'analisis_temporal': self._generate_temporal_analysis(fecha_inicio, fecha_fin),
            'analisis_victimas': self._generate_victim_analysis(),
            'analisis_modalidades': self._generate_modality_analysis(partido),
            'recomendaciones': self._generate_recommendations(partido)
        }
        
        return data
    
    def _format_period(self, fecha_inicio: str, fecha_fin: str) -> str:
        """Formatear el período del informe"""
        inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d')
        fin = datetime.strptime(fecha_fin, '%Y-%m-%d')
        
        if inicio.month == fin.month and inicio.year == fin.year:
            meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
            return f"{meses[inicio.month - 1]} de {inicio.year}"
        else:
            return f"{fecha_inicio} al {fecha_fin}"
    
    def _get_main_crime(self, partido: str) -> str:
        """Obtener el delito principal según el partido"""
        crimes = {
            "José C. Paz": "robo automotor (agravado por el uso de armas de fuego)",
            "San Miguel": "hurto automotor (levantamiento)",
            "Malvinas Argentinas": "robo, en su modalidad de asalto en vía pública",
            "General San Martín": "hurto automotor (levantamiento de autos y motovehículos)"
        }
        return crimes.get(partido, "robo simple")
    
    def _generate_general_stats(self, total: int, partido: str) -> Dict[str, int]:
        """Generar estadísticas generales"""
        # Distribución aproximada basada en los ejemplos
        if partido == "José C. Paz":
            return {
                'homicidios': 3,
                'robos': int(total * 0.65),
                'hurtos': int(total * 0.20),
                'lesiones': int(total * 0.10),
                'otros': int(total * 0.05)
            }
        elif partido == "San Miguel":
            return {
                'homicidios': 2,
                'robos': int(total * 0.55),
                'hurtos': int(total * 0.30),
                'lesiones': int(total * 0.10),
                'otros': int(total * 0.05)
            }
        else:
            return {
                'homicidios': 4,
                'robos': int(total * 0.60),
                'hurtos': int(total * 0.25),
                'lesiones': int(total * 0.10),
                'otros': int(total * 0.05)
            }
    
    def _generate_crime_distribution(self, total: int, partido: str) -> List[Dict[str, Any]]:
        """Generar distribución de delitos por tipo"""
        stats = self._generate_general_stats(total, partido)
        
        distribution = []
        for crime_type, count in stats.items():
            percentage = (count / total) * 100
            distribution.append({
                'tipo': crime_type.title(),
                'cantidad': count,
                'porcentaje': round(percentage, 1)
            })
        
        return sorted(distribution, key=lambda x: x['cantidad'], reverse=True)
    
    def _generate_temporal_analysis(self, fecha_inicio: str, fecha_fin: str) -> Dict[str, Any]:
        """Generar análisis temporal"""
        # Simular datos temporales
        return {
            'dias_mayor_incidencia': ['7', '8', '9'],
            'horarios_criticos': ['18:00-22:00', '02:00-06:00'],
            'tendencia': 'estable',
            'variacion_mensual': '+2.3%'
        }
    
    def _generate_victim_analysis(self) -> Dict[str, Any]:
        """Generar análisis de víctimas"""
        return {
            'genero': {'masculino': 58, 'femenino': 42},
            'edad_promedio': 34,
            'rangos_edad': {
                '18-25': 25,
                '26-35': 30,
                '36-45': 25,
                '46-60': 15,
                '60+': 5
            }
        }
    
    def _generate_modality_analysis(self, partido: str) -> List[Dict[str, Any]]:
        """Generar análisis de modalidades"""
        modalidades = [
            {'modalidad': 'Asalto en vía pública', 'cantidad': 245, 'porcentaje': 32.5},
            {'modalidad': 'Hurto de automotor', 'cantidad': 189, 'porcentaje': 25.1},
            {'modalidad': 'Robo en comercio', 'cantidad': 156, 'porcentaje': 20.7},
            {'modalidad': 'Robo en domicilio', 'cantidad': 98, 'porcentaje': 13.0},
            {'modalidad': 'Otros', 'cantidad': 66, 'porcentaje': 8.7}
        ]
        return modalidades
    
    def _generate_recommendations(self, partido: str) -> List[str]:
        """Generar recomendaciones operativas"""
        return [
            "Intensificar patrullajes en horarios de mayor incidencia (18:00-22:00 hrs)",
            "Reforzar la seguridad en zonas comerciales y vías públicas principales",
            "Implementar operativos preventivos en días de mayor actividad delictiva",
            "Coordinar con fuerzas municipales para control de tránsito vehicular",
            "Fortalecer la presencia policial en transporte público",
            "Desarrollar campañas de prevención dirigidas a la comunidad"
        ]
    
    def _generate_pdf_report(self, data: Dict[str, Any], partido: str, fecha_inicio: str, fecha_fin: str) -> bytes:
        """Generar informe en formato PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
        story = []
        
        # Título principal
        story.append(Paragraph("INFORME SITUACIONAL", self.styles['TituloInforme']))
        story.append(Paragraph(data['periodo'].title(), self.styles['TituloInforme']))
        story.append(Paragraph(partido, self.styles['TituloInforme']))
        story.append(Spacer(1, 0.5*inch))
        
        # Visualizador interactivo (placeholder)
        story.append(Paragraph("VISUALIZADOR INTERACTIVO", self.styles['Subtitulo']))
        story.append(Paragraph("HOMICIDIOS", self.styles['Normal']))
        story.append(Paragraph("ENFRENTAMIENTOS ARMADOS", self.styles['Normal']))
        story.append(Paragraph("Provincia de Buenos Aires", self.styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Introducción
        story.append(Paragraph("INTRODUCCIÓN", self.styles['Subtitulo']))
        intro_text = f"""En {data['periodo']}, entre los delitos y sus modalidades considerados para el presente informe, 
        se registraron {data['total_delitos']} delitos. El {data['delito_principal']} fue el delito más frecuente."""
        story.append(Paragraph(intro_text, self.styles['TextoJustificado']))
        story.append(Spacer(1, 0.2*inch))
        
        # Demografía del partido
        story.append(Paragraph("DEMOGRAFÍA DEL PARTIDO", self.styles['Subtitulo']))
        demo_text = f"""El partido de {partido} presenta características demográficas específicas que influyen 
        en los patrones delictivos observados durante el período analizado."""
        story.append(Paragraph(demo_text, self.styles['TextoJustificado']))
        story.append(Spacer(1, 0.2*inch))
        
        # Estadísticas generales
        story.append(Paragraph("ESTADÍSTICAS GENERALES", self.styles['Subtitulo']))
        
        # Tabla de estadísticas
        stats_data = [['Tipo de Delito', 'Cantidad', 'Porcentaje']]
        for item in data['distribucion_delitos']:
            stats_data.append([item['tipo'], str(item['cantidad']), f"{item['porcentaje']}%"])
        
        stats_table = Table(stats_data, colWidths=[2.5*inch, 1*inch, 1*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Análisis de modalidades
        story.append(Paragraph("ANÁLISIS POR MODALIDADES", self.styles['Subtitulo']))
        modalidades_text = """Las modalidades delictivas muestran patrones específicos que requieren 
        estrategias diferenciadas de prevención y control."""
        story.append(Paragraph(modalidades_text, self.styles['TextoJustificado']))
        
        # Tabla de modalidades
        mod_data = [['Modalidad', 'Cantidad', 'Porcentaje']]
        for item in data['analisis_modalidades']:
            mod_data.append([item['modalidad'], str(item['cantidad']), f"{item['porcentaje']}%"])
        
        mod_table = Table(mod_data, colWidths=[3*inch, 1*inch, 1*inch])
        mod_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(mod_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Análisis temporal
        story.append(Paragraph("ANÁLISIS TEMPORAL", self.styles['Subtitulo']))
        temporal_text = f"""Durante el período analizado, se observó una tendencia {data['analisis_temporal']['tendencia']} 
        en la incidencia delictiva, con mayor concentración en los días {', '.join(data['analisis_temporal']['dias_mayor_incidencia'])} 
        del mes y en los horarios {', '.join(data['analisis_temporal']['horarios_criticos'])}."""
        story.append(Paragraph(temporal_text, self.styles['TextoJustificado']))
        story.append(Spacer(1, 0.2*inch))
        
        # Análisis de víctimas
        story.append(Paragraph("ANÁLISIS DE VÍCTIMAS", self.styles['Subtitulo']))
        victimas_text = f"""El perfil de las víctimas muestra una distribución de {data['analisis_victimas']['genero']['masculino']}% 
        masculino y {data['analisis_victimas']['genero']['femenino']}% femenino, con una edad promedio de 
        {data['analisis_victimas']['edad_promedio']} años."""
        story.append(Paragraph(victimas_text, self.styles['TextoJustificado']))
        story.append(Spacer(1, 0.2*inch))
        
        # Recomendaciones
        story.append(Paragraph("RECOMENDACIONES OPERATIVAS", self.styles['Subtitulo']))
        for i, rec in enumerate(data['recomendaciones'], 1):
            story.append(Paragraph(f"{i}. {rec}", self.styles['Normal']))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Pie de página
        story.append(Paragraph("MAPA GENERAL DE TODAS LAS JURISDICCIONES", self.styles['Subtitulo']))
        mapa_text = """El presente informe se complementa con la visualización cartográfica de los datos, 
        disponible en el sistema de mapas de calor integrado."""
        story.append(Paragraph(mapa_text, self.styles['TextoJustificado']))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _generate_word_report(self, data: Dict[str, Any], partido: str, fecha_inicio: str, fecha_fin: str) -> bytes:
        """Generar informe en formato Word"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Título principal
            title = doc.add_heading('INFORME SITUACIONAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle1 = doc.add_heading(data['periodo'].title(), level=1)
            subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle2 = doc.add_heading(partido, level=1)
            subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Visualizador interactivo
            doc.add_heading('VISUALIZADOR INTERACTIVO', level=2)
            doc.add_paragraph('HOMICIDIOS')
            doc.add_paragraph('ENFRENTAMIENTOS ARMADOS')
            doc.add_paragraph('Provincia de Buenos Aires')
            
            # Introducción
            doc.add_heading('INTRODUCCIÓN', level=2)
            intro_text = f"""En {data['periodo']}, entre los delitos y sus modalidades considerados para el presente informe, 
            se registraron {data['total_delitos']} delitos. El {data['delito_principal']} fue el delito más frecuente."""
            doc.add_paragraph(intro_text)
            
            # Demografía
            doc.add_heading('DEMOGRAFÍA DEL PARTIDO', level=2)
            demo_text = f"""El partido de {partido} presenta características demográficas específicas que influyen 
            en los patrones delictivos observados durante el período analizado."""
            doc.add_paragraph(demo_text)
            
            # Estadísticas generales
            doc.add_heading('ESTADÍSTICAS GENERALES', level=2)
            
            # Tabla de estadísticas
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tipo de Delito'
            hdr_cells[1].text = 'Cantidad'
            hdr_cells[2].text = 'Porcentaje'
            
            for item in data['distribucion_delitos']:
                row_cells = table.add_row().cells
                row_cells[0].text = item['tipo']
                row_cells[1].text = str(item['cantidad'])
                row_cells[2].text = f"{item['porcentaje']}%"
            
            # Recomendaciones
            doc.add_heading('RECOMENDACIONES OPERATIVAS', level=2)
            for i, rec in enumerate(data['recomendaciones'], 1):
                doc.add_paragraph(f"{i}. {rec}")
            
            # Guardar en buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            # Si python-docx no está disponible, generar un PDF
            return self._generate_pdf_report(data, partido, fecha_inicio, fecha_fin)

